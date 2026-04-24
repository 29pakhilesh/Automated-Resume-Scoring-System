"""Extract plain text from resume files (PDF, DOCX).

Includes an OCR fallback for scanned/image-only PDFs when available.
"""

from io import BytesIO
from shutil import which

from docx import Document
from pypdf import PdfReader


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    return "\n".join(parts).strip()


def _tesseract_available() -> bool:
    # pytesseract requires the system `tesseract` binary.
    return which("tesseract") is not None


def extract_text_from_image_ocr(data: bytes) -> str:
    """OCR a single image using Tesseract (best-effort)."""
    if not _tesseract_available():
        raise RuntimeError("OCR unavailable (install tesseract)")
    import pytesseract
    from PIL import Image

    img = Image.open(BytesIO(data)).convert("RGB")
    txt = pytesseract.image_to_string(img)
    return (txt or "").strip()


def extract_text_from_pdf_ocr(data: bytes, max_pages: int = 4) -> str:
    """OCR a PDF by rendering pages to images (pdfium) and running Tesseract.

    - Limits pages for speed; resumes are usually 1–2 pages.
    - Requires: `pypdfium2`, `pillow`, `pytesseract`, and system `tesseract` binary.
    """
    if not _tesseract_available():
        raise RuntimeError("OCR unavailable (install tesseract)")

    import pytesseract
    from PIL import Image
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(data)
    parts: list[str] = []

    # render ~220dpi for decent OCR without being too slow
    scale = 220 / 72

    page_count = min(len(pdf), max_pages)
    for i in range(page_count):
        page = pdf[i]
        bmp = page.render(scale=scale).to_pil()
        img = bmp.convert("RGB") if isinstance(bmp, Image.Image) else bmp
        txt = pytesseract.image_to_string(img)
        if txt:
            parts.append(txt)

    return "\n".join(parts).strip()


def extract_text_from_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text).strip()


def extract_resume_text(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = extract_text_from_pdf(data)
        # If this looks like a scanned/image-only PDF, try OCR (if available).
        if len(text.strip()) < 80:
            try:
                ocr_text = extract_text_from_pdf_ocr(data)
                if len(ocr_text.strip()) > len(text.strip()):
                    return ocr_text
            except Exception:
                # OCR is optional; fall back to normal extraction.
                pass
        return text
    if lower.endswith(".docx"):
        return extract_text_from_docx(data)
    if lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
        # Always OCR for image resumes.
        return extract_text_from_image_ocr(data)
    raise ValueError("Unsupported file type. Use PDF, DOCX, PNG, JPG, or WEBP.")
